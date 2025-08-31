"""
RAG Knowledge Base System - Native Vector Store Implementation

This module provides a native implementation of a vector store system without external
dependencies like Langchain. It handles document processing, chunking, embedding creation,
and storage using OpenAI's API directly.

Features:
- Native document loading from multiple formats (PDF, DOCX, Excel, TXT, JSON)
- Token-based text chunking using tiktoken
- Batch processing for efficient embedding creation
- JSON-based vector store storage
- Comprehensive metadata preservation
- Direct OpenAI API integration

Author: Usuf Com
Contact: usufcom20@gmail.com
Website: www.djamai.com
LinkedIn: https://www.linkedin.com/in/usufcom
Date: 2024
"""

import os
import json
import numpy as np
import tiktoken
from openai import OpenAI
from typing import List, Dict
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
import xlrd


class VectorStore:
    """
    A native vector store implementation for RAG systems.
    
    This class provides a complete vector store solution without external dependencies,
    handling document processing, chunking, embedding creation, and storage.
    
    Attributes:
        client (OpenAI): OpenAI client for API calls
        chunk_size (int): Maximum number of tokens per chunk
        chunk_overlap (int): Number of tokens to overlap between chunks
        batch_size (int): Number of texts to process in each embedding batch
    """
    
    def __init__(self, api_key: str, chunk_size: int = 1000, chunk_overlap: int = 200, batch_size: int = 16):
        """
        Initialize the native vector store system.
        
        Args:
            api_key (str): OpenAI API key for embeddings
            chunk_size (int): Maximum tokens per text chunk (default: 1000)
            chunk_overlap (int): Token overlap between chunks (default: 200)
            batch_size (int): Number of texts to embed in each batch (default: 16)
        
        Note:
            chunk_overlap should be less than chunk_size for meaningful chunks.
            batch_size affects API call efficiency - larger batches are more efficient
            but may hit rate limits with very large documents.
        """
        # Initialize OpenAI client for API calls
        self.client = OpenAI(api_key=api_key)
        
        # Store chunking configuration
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.batch_size = batch_size

    def load_documents(self, folder_path: str) -> List[Dict]:
        """
        Load documents from a folder and convert them to a standardized format.
        
        This method handles multiple file formats and extracts text content while
        preserving metadata about the source file, page numbers, and content type.
        
        Supported formats:
        - PDF files: Extracts text from each page with page numbers
        - Word documents: Extracts paragraphs and tables separately
        - Excel files: Converts each sheet to text with sheet information
        - Text files: Direct text extraction
        - JSON files: Converts to string representation
        
        Args:
            folder_path (str): Path to the folder containing documents
        
        Returns:
            List[Dict]: List of documents with 'page_content' and 'metadata' keys
        
        Note:
            Each document includes metadata with source filename, content type,
            and additional context (page numbers, sheet names, etc.).
        """
        docs = []
        
        # Process files in sorted order for consistent results
        for filename in sorted(os.listdir(folder_path)):
            file_path = os.path.join(folder_path, filename)
            
            # Skip directories and non-files
            if not os.path.isfile(file_path):
                continue

            # Convert filename to lowercase for case-insensitive matching
            lower = filename.lower()
            
            try:
                # Handle PDF files - extract text from each page
                if lower.endswith(".pdf"):
                    reader = PdfReader(file_path)
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text() or ""
                        if text.strip():  # Only add non-empty pages
                            docs.append({
                                "page_content": text,
                                "metadata": {
                                    "source": filename, 
                                    "page": i + 1, 
                                    "type": "pdf"
                                }
                            })

                # Handle Word documents - extract paragraphs and tables
                elif lower.endswith(".docx"):
                    doc = DocxDocument(file_path)
                    
                    # Extract paragraphs
                    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
                    if paragraphs:
                        docs.append({
                            "page_content": "\n".join(paragraphs),
                            "metadata": {
                                "source": filename, 
                                "type": "docx_paragraphs"
                            }
                        })
                    
                    # Extract tables separately for better structure
                    for t_idx, table in enumerate(doc.tables):
                        rows = []
                        for row in table.rows:
                            # Extract cell text and join with pipe separator
                            cells = [cell.text.strip() if cell.text else "" for cell in row.cells]
                            rows.append(" | ".join(cells))
                        
                        table_text = "\n".join(rows).strip()
                        if table_text:  # Only add non-empty tables
                            docs.append({
                                "page_content": table_text,
                                "metadata": {
                                    "source": filename, 
                                    "table_index": t_idx, 
                                    "type": "docx_table"
                                }
                            })

                # Handle plain text files
                elif lower.endswith(".txt"):
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    if text.strip():  # Only add non-empty files
                        docs.append({
                            "page_content": text, 
                            "metadata": {
                                "source": filename, 
                                "type": "txt"
                            }
                        })

                # Handle JSON files - convert to string representation
                elif lower.endswith(".json"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    docs.append({
                        "page_content": json.dumps(data, ensure_ascii=False),
                        "metadata": {
                            "source": filename, 
                            "type": "json"
                        }
                    })

                # Handle modern Excel files (.xlsx)
                elif lower.endswith(".xlsx"):
                    wb = load_workbook(file_path, read_only=True, data_only=True)
                    for sheet in wb.worksheets:
                        rows_text = []
                        for row in sheet.iter_rows(values_only=True):
                            # Convert row values to strings and filter out empty cells
                            row_vals = [str(c) for c in row if c]
                            if row_vals:  # Only add non-empty rows
                                rows_text.append(" | ".join(row_vals))
                        
                        sheet_text = "\n".join(rows_text).strip()
                        if sheet_text:  # Only add non-empty sheets
                            docs.append({
                                "page_content": sheet_text,
                                "metadata": {
                                    "source": filename, 
                                    "sheet": sheet.title, 
                                    "type": "xlsx"
                                }
                            })

                # Handle legacy Excel files (.xls)
                elif lower.endswith(".xls"):
                    book = xlrd.open_workbook(file_path)
                    for si in range(book.nsheets):
                        sheet = book.sheet_by_index(si)
                        rows_text = []
                        
                        for ri in range(sheet.nrows):
                            # Extract row values and filter out empty cells
                            row_vals = [str(c) for c in sheet.row_values(ri) if c]
                            if row_vals:  # Only add non-empty rows
                                rows_text.append(" | ".join(row_vals))
                        
                        text = "\n".join(rows_text).strip()
                        if text:  # Only add non-empty sheets
                            docs.append({
                                "page_content": text,
                                "metadata": {
                                    "source": filename, 
                                    "sheet_index": si, 
                                    "type": "xls"
                                }
                            })

            except Exception as e:
                # Log errors but continue processing other files
                print(f"⚠️ Error loading {filename}: {e}")
                continue
                
        return docs

    def chunk_text(self, doc: dict) -> List[Dict]:
        """
        Split a document's text into overlapping chunks using token-based splitting.
        
        This method uses tiktoken to split text by tokens rather than characters,
        which provides more consistent chunk sizes across different languages and
        text types. Each chunk preserves the original document's metadata.
        
        Args:
            doc (dict): Document dictionary with 'page_content' and 'metadata' keys
        
        Returns:
            List[Dict]: List of chunks, each with 'text' and 'metadata' keys
        
        Note:
            Uses the cl100k_base encoding which is compatible with GPT models.
            Chunk overlap helps maintain context between chunks.
        """
        # Ensure text is a string
        text = str(doc["page_content"])
        metadata = doc.get("metadata", {})

        # Get tiktoken encoder for token-based splitting
        enc = tiktoken.get_encoding("cl100k_base")
        tokens = enc.encode(text)

        chunks = []
        start = 0
        
        # Create overlapping chunks
        while start < len(tokens):
            # Calculate chunk end position
            end = min(start + self.chunk_size, len(tokens))
            
            # Decode tokens back to text
            chunk = enc.decode(tokens[start:end])
            
            # Create chunk with preserved metadata
            chunks.append({
                "text": chunk,
                "metadata": metadata.copy()  # Copy metadata to avoid reference issues
            })
            
            # Move start position for next chunk (with overlap)
            start += self.chunk_size - self.chunk_overlap

        return chunks

    def embed_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """
        Create embeddings for text chunks in batches for efficiency.
        
        This method processes chunks in batches to minimize API calls and improve
        performance. Each chunk's metadata is preserved in the output.
        
        Args:
            chunks (List[Dict]): List of chunks with 'text' and 'metadata' keys
        
        Returns:
            List[Dict]: List of embeddings with 'text', 'embedding', and 'metadata' keys
        
        Note:
            Uses text-embedding-3-small model for optimal performance and cost.
            Batch processing reduces API overhead and improves throughput.
        """
        embeddings = []
        
        # Process chunks in batches
        for i in range(0, len(chunks), self.batch_size):
            batch = chunks[i:i + self.batch_size]
            texts = [c["text"] for c in batch]

            # Create embeddings for the batch
            response = self.client.embeddings.create(
                model="text-embedding-3-small",
                input=texts
            )

            # Combine embeddings with original chunk data
            for j, emb in enumerate(response.data):
                embeddings.append({
                    "text": batch[j]["text"],
                    "embedding": emb.embedding,
                    "metadata": batch[j]["metadata"]  # Preserve original metadata
                })
                
        return embeddings

    def save_vector_store_local(self, embeddings: List[Dict], vector_store_path: str):
        """
        Save embeddings to a local JSON file.
        
        This method creates the necessary directories and saves the vector store
        as a JSON file for later loading and querying.
        
        Args:
            embeddings (List[Dict]): List of embeddings with text, embedding, and metadata
            vector_store_path (str): Path where the JSON file should be saved
        
        Note:
            The JSON format preserves all metadata and is human-readable.
            Uses UTF-8 encoding to handle international characters.
        """
        # Ensure the directory exists
        os.makedirs(os.path.dirname(vector_store_path), exist_ok=True)

        # Save embeddings as JSON with proper formatting
        try:
            with open(vector_store_path, "w", encoding="utf-8") as json_file:
                json.dump(embeddings, json_file, ensure_ascii=False, indent=2)
            print(f"✅ Vector store saved to {vector_store_path}")
        except PermissionError:
            print(f"❌ Permission denied. Make sure {vector_store_path} is not open in another program.")

    def exract_save_vector_store(self, store, kb_folder: str, vector_store_path: str):
        """
        Complete pipeline to extract, process, and save a vector store.
        
        This method orchestrates the entire process:
        1. Load documents from the specified folder
        2. Chunk each document into smaller pieces
        3. Create embeddings for all chunks
        4. Save the complete vector store to disk
        
        Args:
            store: Self-reference for method chaining
            kb_folder (str): Path to the knowledge base folder
            vector_store_path (str): Path where the vector store should be saved
        
        Note:
            This is the main entry point for creating a new vector store.
            The process is memory-efficient as it processes documents sequentially.
        """
        # Step 1: Load all documents from the knowledge base folder
        docs = self.load_documents(kb_folder)

        # Step 2: Process each document - chunk and embed
        all_embeddings = []
        for doc in docs:  # Each doc is {"page_content": ..., "metadata": {...}}
            # Split document into chunks while preserving metadata
            chunks = store.chunk_text(doc)
            
            # Create embeddings for the chunks
            doc_embeddings = store.embed_chunks(chunks)
            
            # Add to the complete collection
            all_embeddings.extend(doc_embeddings)

        # Step 3: Save the complete vector store locally
        self.save_vector_store_local(all_embeddings, vector_store_path)

    